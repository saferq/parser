import os
from docx.enum.text import WD_COLOR_INDEX
from docx import Document
import pandas as pd
from pprint import pprint as pp

import copy
import re


def get_df_excel():
    df = pd.read_excel('ntd.xlsx', sheet_name='NTD', header=0)

    arr_df = []
    for index, row in df.iterrows():
        arr_df.append(list(row))
    # pp(arr_df)
    return arr_df


def replace_docx(arr, path, names_files):

    for name_file in names_files:
        print(name_file)
        doc = Document(f'{path}\\{name_file}')
        count = 0
        for paragraph in doc.paragraphs:
            for target in arr:
                ntd_short = target[0]
                ntd_full = target[1]
                ntd_search = target[2]
                if ntd_search in paragraph.text:
                    # print("\n\tparagraph.text")
                    # print(paragraph.text)

                    change_paragraph = paragraph.text
                    paragraph.text = ''

                    change_paragraph = re.split(
                        f'({ntd_search})', change_paragraph)

                    # print(change_paragraph)

                    for text in change_paragraph:
                        count += 1
                        if ntd_search == text:
                            print(f'<<<<: {text} < {ntd_short}')
                            newRun = paragraph.add_run(text)
                            newRun.font.highlight_color = WD_COLOR_INDEX.PINK
                            addRun = paragraph.add_run(
                                f' {ntd_full}')
                            addRun.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            paragraph.add_run(text)

                   # # deep copy as we delete/clear the object
                    # # currRuns = copy.copy(paragraph.runs)
                    # currRuns = paragraph.runs
                    # paragraph.runs.clear()

                    # for run in currRuns:
                    #     print(f'runs: {run.text}')
                    #     if target[2] in run.text:

                    #         # split into words in order to be able to color only one
                    #         words = re.split('(\s)', run.text)
                    #         # print(words)
                    #         for word in words:
                    #             if target[2] in word:
                    #                 print(f'<<<<: {word} < {target[0]}')
                    #                 newRun = paragraph.add_run(word)
                    #                 newRun.font.highlight_color = WD_COLOR_INDEX.PINK
                    #                 addRun = paragraph.add_run(
                    #                     f' [заменен на {target[1]}] ')
                    #                 addRun.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    #                 count += 1
                    #             else:
                    #                 newRun = paragraph.add_run(word)
                    #                 newRun.font.highlight_color = None
                    #     else:  # our target is not in it so we add it unchanged
                    #         paragraph.runs.append(run)

        if count > 0:
            print(f'save: {name_file}\n')
            doc.save(f'{name_file}')
        else:
            print('-----\n')


def get_names_files(path):
    names_files = (os.listdir(path))
    # print(names_files)
    if names_files == []:
        print(f'Папка "{path}" пустая')
    return path, names_files


if __name__ in '__main__':

    arr = get_df_excel()
    path, names_files = get_names_files('files')
    replace_docx(arr, path, names_files)
